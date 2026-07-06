"""Parsing + segmentation tests (offline, no TTS model)."""

from __future__ import annotations

import io

import docx
import pytest

from sangyin.models import Sentence
from sangyin.parsing import parse_text, parse_upload
from sangyin.parsing.base import build_document, group_sentences, segment_sentences


def test_segment_sentences_basic():
    assert segment_sentences("One. Two! Three?") == ["One.", "Two!", "Three?"]
    assert segment_sentences("   ") == []


def test_group_sentences_merges_short_and_splits_long():
    # Short sentences collapse into one phrase; every sentence is kept, in order.
    short = [Sentence(index=i, text=t) for i, t in enumerate(["One.", "Two.", "Three."])]
    groups = group_sentences(short)
    assert [s.index for g in groups for s in g] == [0, 1, 2]
    assert len(groups) == 1

    # Long sentences split across phrases; the small first group starts playback fast.
    long = [Sentence(index=i, text="word " * 30) for i in range(4)]
    groups = group_sentences(long)
    assert len(groups) > 1
    assert groups[0][0].index == 0
    # No sentence is dropped or duplicated.
    assert [s.index for g in groups for s in g] == [0, 1, 2, 3]
    # The max_count cap is respected.
    assert all(len(g) <= 4 for g in groups)


def test_pasted_text_global_indices():
    doc = parse_text("Alpha one. Beta two. Gamma three.", title="Demo")
    sentences = [s for c in doc.chapters for s in c.sentences]
    assert [s.index for s in sentences] == [0, 1, 2]
    assert sentences[0].text == "Alpha one."
    assert doc.source_type == "text"
    assert doc.n_sentences == 3


def test_build_document_indices_span_chapters():
    doc = build_document(
        title="Multi",
        source_type="txt",
        raw_chapters=[("A", "One. Two."), ("B", "Three. Four. Five.")],
    )
    assert len(doc.chapters) == 2
    indices = [s.index for c in doc.chapters for s in c.sentences]
    # Indices are global and contiguous across chapters.
    assert indices == [0, 1, 2, 3, 4]
    assert doc.chapters[1].sentences[0].index == 2


def test_docx_splits_on_headings():
    d = docx.Document()
    d.add_heading("Chapter One", level=1)
    d.add_paragraph("First sentence here. Second sentence here.")
    d.add_heading("Chapter Two", level=1)
    d.add_paragraph("Third sentence here. Fourth sentence here.")
    buf = io.BytesIO()
    d.save(buf)

    doc = parse_upload("book.docx", buf.getvalue())
    titles = [c.title for c in doc.chapters]
    assert titles == ["Chapter One", "Chapter Two"]
    assert doc.source_type == "docx"


def test_epub_chapters(tmp_path):
    ebooklib = pytest.importorskip("ebooklib")
    from ebooklib import epub

    book = epub.EpubBook()
    book.set_title("My Book")
    c1 = epub.EpubHtml(title="Ch1", file_name="c1.xhtml")
    c1.content = "<h1>Chapter One</h1><p>Alpha sentence. Beta sentence.</p>"
    c2 = epub.EpubHtml(title="Ch2", file_name="c2.xhtml")
    c2.content = "<h1>Chapter Two</h1><p>Gamma sentence. Delta sentence.</p>"
    book.add_item(c1)
    book.add_item(c2)
    book.toc = [c1, c2]
    book.spine = [c1, c2]
    book.add_item(epub.EpubNcx())
    book.add_item(epub.EpubNav())

    path = tmp_path / "book.epub"
    epub.write_epub(str(path), book)

    doc = parse_upload("book.epub", path.read_bytes())
    assert doc.source_type == "epub"
    titles = [c.title for c in doc.chapters]
    assert "Chapter One" in titles and "Chapter Two" in titles


def test_pdf_outline_splits_chapters():
    reportlab = pytest.importorskip("reportlab")
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    c.drawString(72, 720, "Front matter introduction text.")
    c.showPage()
    c.bookmarkPage("ch1")
    c.addOutlineEntry("Chapter One", "ch1", level=0)
    c.drawString(72, 720, "Chapter one body sentence.")
    c.showPage()
    c.bookmarkPage("ch2")
    c.addOutlineEntry("Chapter Two", "ch2", level=0)
    c.drawString(72, 720, "Chapter two body sentence.")
    c.showPage()
    c.save()

    doc = parse_upload("doc.pdf", buf.getvalue())
    assert doc.source_type == "pdf"
    titles = [c.title for c in doc.chapters]
    assert "Chapter One" in titles and "Chapter Two" in titles
    # Front matter before the first bookmark becomes its own chapter.
    assert "Front matter" in titles


def test_unsupported_extension_raises():
    with pytest.raises(ValueError):
        parse_upload("weird.xyz", b"data")
